from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent / "meeting-briefs" / "Alem_FeSMA_Processing_Brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
HEADER_FILL = "F2F4F7"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_metadata(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=BLACK, bold=True)
    run = p.add_run(value)
    set_run_font(run, size=10.5, color=BLACK)


def add_bottom_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A6A6A6")
    borders.append(bottom)
    p_pr.append(borders)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_compact_table(doc):
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run("Table 1. Recorded processing routes used for the internal comparison.")
    set_run_font(run, size=10, color=GRAY, italic=True)

    rows = [
        ("Step", "Candidate Fe-SMA", "Benchmark Fe-SMA"),
        ("Melt and cast", "Vacuum induction melt at 1650 °C for 30 min; poured into a water-cooled copper mould.", "Vacuum induction melt; remelted on an Arcast unit after voiding was observed; cast as a 0.600 in × 6 in rod."),
        ("Stock and homogenization", "Two 0.500 in rods were cut by EDM. Tested wire was from the 1000 °C / 16 h argon-homogenized route.", "The benchmark was also confirmed to receive the 1000 °C / 16 h argon homogenization treatment."),
        ("Hot deformation", "Hot rolled at 850 °C to approximately 0.175 in square section.", "Hot swaged at 900 °C to 0.475 in, then hot rolled at 900 °C to 0.210 in square section."),
        ("Cold work and process annealing", "Cold swaged/drawn with 1000 °C argon process anneals to restore ductility; final wire 0.0142 in (0.36 mm).", "Cold swaged and drawn with a 900 °C / 1 min anneal, followed by 1000 °C process anneals in H₂ (F72) and argon; final wire 0.0253 in (0.64 mm)."),
        ("Shared final treatment", "Both alloys: sealed in quartz under argon; 1200 °C / 30 min cycles with 10 min cold-zone transfers, followed by water quench.", "Both alloys: sealed in quartz under argon; 1200 °C / 30 min cycles with 10 min cold-zone transfers, followed by water quench."),
    ]
    table = doc.add_table(rows=0, cols=3)
    configure_table(table, [1940, 3710, 3710])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=BLACK, bold=(row_idx == 0 or idx == 0))
            if row_idx == 0:
                set_cell_shading(cells[idx], HEADER_FILL)
    return table


def main():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Alem Meeting Brief | Experimental Processing")
    set_run_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Internal working summary | 26 August 2026")
    set_run_font(run, size=8.5, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Fe-SMA Experimental Processing Summary")
    set_run_font(run, size=23, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Procedures-focused briefing for Alem")
    set_run_font(run, size=13, color=GRAY)

    add_metadata(doc, "Purpose", "Review the melting, hot-working, drawing, and heat-treatment routes used for the candidate and benchmark Fe-SMA wires.")
    add_metadata(doc, "Scope", "Experimental processing only; thermodynamic calculations, phase analysis, microscopy, and synchrotron measurements are intentionally excluded.")
    add_bottom_rule(doc)

    doc.add_paragraph("1. Materials and processing objective", style="Heading 1")
    add_body(doc, "The candidate Fe-SMA was designed with a nominal composition of 50Fe-30Mn-12Al-4Ni-4Si (at%) with 1000 ppm C. The measured wire chemistry was 51.5Fe-29.8Mn-11.9Al-4.2Ni-2.0Si-0.45C (at%). A benchmark Fe-SMA, nominally 43.5Fe-34Mn-15Al-7.5Ni (at%), was processed in parallel as a reference route. Both routes were taken to fine wire through high-temperature deformation, cold work, and intermediate annealing.")

    doc.add_paragraph("2. Route comparison", style="Heading 1")
    add_compact_table(doc)

    doc.add_paragraph("3. Candidate-alloy route in sequence", style="Heading 1")
    add_body(doc, "High-purity charge materials were vacuum-induction melted at 1650 °C for 30 min and poured into a water-cooled copper mould. Two 0.500 in rods were cut from the casting by EDM. The tested candidate wire came from the homogenized route: 1000 °C for 16 h in argon.")
    add_body(doc, "The homogenized stock was hot rolled at 850 °C in multiple passes to approximately 0.175 in square. It was then cold swaged and cold drawn. Work hardening during drawing was managed with 1000 °C argon process anneals, after which drawing continued to a final diameter of 0.0142 in (0.36 mm). The final cold-drawing reduction after the last process anneal was approximately 85% in area.")

    doc.add_paragraph("4. Meeting checks", style="Heading 1")
    add_body(doc, "The final tested candidate wire is confirmed to be from the homogenized route. The side-by-side details above retain the process-log distinctions between the two materials, especially the 850 °C versus 900 °C hot-working temperatures, different final wire diameters, and benchmark use of hydrogen during part of the drawing route. These are the most useful items to confirm or refine during the next processing review.")

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(8)
    source.paragraph_format.space_after = Pt(0)
    run = source.add_run("Source basis: process-note-SCai; PROCESSING-AND-REPLICATES; current Fe-SMA manuscript experimental section.")
    set_run_font(run, size=9, color=GRAY, italic=True)

    doc.core_properties.title = "Fe-SMA Experimental Processing Summary"
    doc.core_properties.subject = "Procedures-focused meeting brief for Alem"
    doc.core_properties.author = "Frank Y. Cai"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
